from docx import Document

# Creating the document
doc = Document()
doc.add_heading('Dicionário de Dados - Produção Agrícola', level=1)

# Adding table descriptions
tables = [
    {
        "name": "Producao",
        "description": "Armazena informações sobre a produção agrícola, incluindo estimativas e quantidades efetivas para diversas culturas e regiões.",
        "fields": [
            {"name": "id", "type": "NUMBER", "description": "Identificador único da produção (auto-incrementado)", "detail": "Chave Primária"},
            {"name": "id_produto", "type": "NUMBER", "description": "Identificador do produto cultivado", "detail": "Chave Estrangeira para a tabela Produto"},
            {"name": "id_uf", "type": "VARCHAR2(2)", "description": "Código do estado onde ocorre a produção", "detail": "Chave Estrangeira para a tabela UF"},
            {"name": "id_safra", "type": "NUMBER", "description": "Identificador do período de safra", "detail": "Chave Estrangeira para a tabela Safra"},
            {"name": "area_plantada", "type": "NUMBER", "description": "Área total plantada"},
            {"name": "qtde_estimada", "type": "NUMBER", "description": "Quantidade estimada de produção"},
            {"name": "qtde_efetiva", "type": "NUMBER", "description": "Quantidade efetiva de produção"},
        ]
    },
    {
        "name": "Produto",
        "description": "Armazena informações sobre os tipos de produtos agrícolas.",
        "fields": [
            {"name": "id", "type": "NUMBER", "description": "Identificador único do produto.", "detail": "Chave Primária"},
            {"name": "nome", "type": "VARCHAR2(255)", "description": "Nome do produto."},
        ]
    },
    {
        "name": "UF",
        "description": "Armazena dados sobre as unidades federativas (estados) do Brasil.",
        "fields": [
            {"name": "id", "type": "VARCHAR2(2)", "description": "Código do estado.", "detail": "Chave Primária"},
            {"name": "id_regiao", "type": "NUMBER", "description": "Identificador da região.", "detail": "Chave Estrangeira para a tabela Regiao"},
            {"name": "nome", "type": "VARCHAR2(255)", "description": "Nome do estado."},
        ]
    },
    {
        "name": "Safra",
        "description": "Armazena informações sobre períodos de safra, especificando ano e mês.",
        "fields": [
            {"name": "id", "type": "NUMBER", "description": "Identificador único da safra.", "detail": "Chave Primária"},
            {"name": "ano", "type": "NUMBER", "description": "Ano da safra."},
            {"name": "mes", "type": "NUMBER", "description": "Mês da safra."},
        ]
    },
    {
        "name": "Regiao",
        "description": "Armazena as regiões do Brasil, agrupando unidades federativas.",
        "fields": [
            {"name": "id", "type": "NUMBER", "description": "Identificador único da região.", "detail": "Chave Primária"},
            {"name": "nome", "type": "VARCHAR2(255)", "description": "Nome da região."},
        ]
    }
]

# Adding each table to the document
for table in tables:
    doc.add_heading(f"Tabela: {table['name']}", level=2)
    doc.add_paragraph(f"Descrição: {table['description']}")
    doc.add_paragraph("Campos:", style="Heading 3")

    table_obj = doc.add_table(rows=1, cols=4)
    hdr_cells = table_obj.rows[0].cells
    hdr_cells[0].text = 'Nome'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Descrição'
    hdr_cells[3].text = 'Detalhe'

    for field in table["fields"]:
        row_cells = table_obj.add_row().cells
        row_cells[0].text = field["name"]
        row_cells[1].text = field["type"]
        row_cells[2].text = field["description"]
        row_cells[3].text = field.get("detail", "")

    doc.add_paragraph("")

# Save the document
doc.save("Dicionario_de_Dados_Producao_Agricola.docx")
